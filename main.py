from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import yaml
import subprocess
import base64
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI()

class ResumeRequest(BaseModel):
    avid_yaml: str
    aifa_yaml: str
    studio_yaml: str
    aiesec_yaml: str
    skills_latex: str
    summary_latex: str
    output_format: str = "pdf"

def unescape(text: str) -> str:
    return text.replace('\\n', '\n')

def parse_role(yaml_str: str) -> dict:
    return yaml.safe_load(unescape(yaml_str))

def build_rendercv_yaml(roles: list) -> dict:
    work_experience = []
    for role in roles:
        entry = {
            'company': role.get('company', ''),
            'position': role.get('role', ''),
            'date': role.get('date_range', ''),
            'highlights': role.get('bullets', [])
        }
        work_experience.append(entry)

    return {
        'cv': {
            'name': 'Apeksha Reddy Saddi, PMP, CSM',
            'location': 'Seattle, WA, USA',
            'email': 'saddi.apeksha@gmail.com',
            'phone': '+1 609-255-9917',
            'social_networks': [
                {'network': 'LinkedIn', 'username': 'apeksha-saddi'}
            ],
            'sections': {
                'work_experience': work_experience
            }
        },
        'design': {
            'theme': 'classic'
        }
    }

def build_docx(roles: list) -> bytes:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('Apeksha Reddy Saddi, PMP, CSM')
    name_run.bold = True
    name_run.font.size = Pt(14)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run(
        'saddi.apeksha@gmail.com  |  +1 609-255-9917  |  '
        'Seattle, WA, USA (Open to Relocate)  |  '
        'linkedin.com/in/apeksha-saddi'
    )

    doc.add_heading('WORK EXPERIENCE', level=1)

    for role in roles:
        title = role.get('role', '')
        company = role.get('company', '')
        date_range = role.get('date_range', '')
        bullets = role.get('bullets', [])

        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f'{title} - {company}')
        role_run.bold = True
        role_para.add_run(f'  {date_range}')

        for bullet in bullets:
            bullet_para = doc.add_paragraph(
                bullet, style='List Bullet'
            )
            bullet_para.paragraph_format.space_after = Pt(2)

    with tempfile.NamedTemporaryFile(
        suffix='.docx', delete=False
    ) as f:
        doc.save(f.name)
        with open(f.name, 'rb') as docx_file:
            content = docx_file.read()
    os.unlink(f.name)
    return content

def render_pdf_with_rendercv(roles: list) -> bytes:
    cv_data = build_rendercv_yaml(roles)

    with tempfile.TemporaryDirectory() as tmpdir:
        yaml_path = os.path.join(tmpdir, 'cv.yaml')
        with open(yaml_path, 'w') as f:
            yaml.dump(cv_data, f, allow_unicode=True)

        result = subprocess.run(
            ['python', '-m', 'rendercv', 'render', yaml_path],
            capture_output=True,
            text=True,
            cwd=tmpdir
        )

        if result.returncode != 0:
            raise HTTPException(
                status_code=500,
                detail=f"RenderCV error: {result.stderr}"
            )

        output_dir = os.path.join(
            tmpdir,
            'rendercv_output'
        )

        pdf_files = []
        for root, dirs, files in os.walk(tmpdir):
            for file in files:
                if file.endswith('.pdf'):
                    pdf_files.append(os.path.join(root, file))

        if not pdf_files:
            raise HTTPException(
                status_code=500,
                detail=f"No PDF found. stdout: {result.stdout} stderr: {result.stderr}"
            )

        with open(pdf_files[0], 'rb') as f:
            return f.read()

@app.post("/render")
async def render_resume(request: ResumeRequest):
    try:
        roles = [
            parse_role(request.avid_yaml),
            parse_role(request.aifa_yaml),
            parse_role(request.studio_yaml),
            parse_role(request.aiesec_yaml),
        ]
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"YAML parse error: {str(e)}"
        )

    if request.output_format == "docx":
        try:
            docx_bytes = build_docx(roles)
            return {
                "format": "docx",
                "content": base64.b64encode(docx_bytes).decode()
            }
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"DOCX error: {str(e)}"
            )
    else:
        try:
            pdf_bytes = render_pdf_with_rendercv(roles)
            return {
                "format": "pdf",
                "content": base64.b64encode(pdf_bytes).decode()
            }
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"PDF error: {str(e)}"
            )

@app.get("/health")
async def health():
    return {"status": "ok"}
